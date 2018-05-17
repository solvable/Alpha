import datetime
import os
from io import BytesIO

from django.http import HttpResponse


from django.shortcuts import render
from django.views import generic
from .models import Estimate, Section
from CRM.models import Customer, Jobsite, Ticket
from django.shortcuts import get_object_or_404
from django.forms import inlineformset_factory
from django.views import generic
from django.shortcuts import redirect
from .forms import SectionFormset, EstimateForm, SectionInline, SectionForm
from django.urls import reverse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import datetime









class EstimateCreateView(generic.CreateView):
    model = Estimate
    template_name = 'estimate/create_estimate.html'
    fields = ['customer', 'jobsite', 'ticket', 'name', 'billStreet', 'billCityStateZip', 'phone', 'email', 'job_address']

    def get_success_url(self):
        return reverse('estimate-detail', kwargs={'cust': self.object.customer, 'job': self.object.jobsite, 'ticket': self.object.ticket, 'est': self.object.id})

    def get_initial(self):
        customer = get_object_or_404(Customer, id=self.kwargs.get('cust'))
        jobsite = get_object_or_404(Jobsite, id=self.kwargs.get('job'))
        ticket =  get_object_or_404(Ticket, id=self.kwargs.get('ticket'))

        return{
            'customer':customer.id,
            'jobsite':jobsite.id,
            'ticket':ticket.id,
            'name': customer.fullName,
            'billStreet': customer.billStreet,
            'billCityStateZip': "%s, %s, %s" % (customer.billCity, customer.billState, customer.billZip),
            'job_address': jobsite.jobStreet,
            'phone': customer.phone1,
            'email':customer.email,
        }

    def get_context_data(self, **kwargs):
        context = super(EstimateCreateView, self).get_context_data(**kwargs)

        if self.request.POST:
            context['sectionform'] = SectionFormset(self.request.POST, prefix='section')
        else:
            context['sectionform'] = SectionFormset(prefix='section')
        return context


    def form_valid(self, form):
        context = self.get_context_data()
        section_form = context['sectionform']

        if form.is_valid() and section_form.is_valid():
            self.object = form.save(commit=False)
            section_form.instance = self.object
            section_form.save(commit=False)
            form.save()
            section_form.save()
            form.save()




            return redirect(self.get_success_url())
        else:
            return self.render_to_response(self.get_context_data(form=form))

class EstimateUpdateView(generic.UpdateView):
    model = Estimate
    template_name = 'estimate/create_estimate.html'
    pk_url_kwarg = 'est'
    fields='__all__'

    def get_success_url(self):
        return reverse('estimate-detail', kwargs={'cust': self.object.customer, 'job': self.object.jobsite, 'ticket': self.object.ticket, 'est': self.object.id})


    def get_context_data(self, **kwargs):
        context = super(EstimateUpdateView, self).get_context_data(**kwargs)

        if self.request.POST:
            context['sectionform'] = SectionFormset(self.request.POST, instance=self.object, prefix='section')

        else:
            context['sectionform'] = SectionFormset(instance=self.object, prefix='section')

        return context

    def form_valid(self, form):
        context = self.get_context_data()
        section_form = context['sectionform']

        if form.is_valid() and section_form.is_valid():
            self.object = form.save(commit=False)
            section_form.instance = self.object
            section_form.save(commit=False)
            section_form.save()
            form.save()
            return redirect(self.get_success_url())
        else:
            return self.render_to_response(self.get_context_data(form=form))







class EstimateDetailView(generic.DetailView):
    model=Estimate
    template_name = 'estimate/estimate_detail.html'
    pk_url_kwarg = "est"



class EstimateDeleteView(generic.DeleteView):
    model = Estimate
    template_name = 'estimate/estimate_confirm_delete.html'
    pk_url_kwarg = "est"

    def get_success_url(self):
        return reverse('ticket_detail', kwargs={'cust': self.kwargs.get('cust'), 'job': self.kwargs.get('job'), 'ticket': self.kwargs.get('ticket')})



def write_invoice_view(request, cust, job, ticket, est):
    import html2text
    import re
    from docx.shared import RGBColor, Pt, Inches
    from docx.enum.text import WD_COLOR_INDEX

    current_path = request.get_full_path()
    print(current_path)

    # load objects
    customer = get_object_or_404(Customer, id=cust)
    jobsite = get_object_or_404(Jobsite, id=job)
    ticket = get_object_or_404(Ticket, id=ticket)
    estimate = get_object_or_404(Estimate, id=est)

    estimate.completed = True
    estimate.save()

    # setup buffer
    buffer = BytesIO()

    # Set Some Variables for filename and path
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    # headerimage = os.path.join(BASE_DIR, "static/CRM/img/logo-redux.png")
    filename = customer.lastName + "_" + customer.firstName + "-" + jobsite.jobStreet + "-workorder#" + str(ticket.id)
    doc = os.path.join(BASE_DIR, "static/CRM/doc-template/newest.docx")

    # Set misc variables
    blank_space = 6
    topmargin = .50
    gutters = .50
    bottommargin = .25

    # Create HttpResponse object with appropriate PDF headers
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachement; filename=' + filename + ".docx"

    # Create the docx object
    document = Document(doc)




    sections = document.sections
    for section in sections:
        section.top_margin = Inches(topmargin)
        section.bottom_margin = Inches(bottommargin)
        section.left_margin = Inches(gutters)
        section.right_margin = Inches(gutters)



    # Set the filename variable
    docx_title = "%s.docx" % (filename,)

    # Add content to docx file


    # setup variables for docx table
    today = datetime.date.today()
    bill1 = customer.billStreet
    bill2 = '%s, %s, %s' % (customer.billCity, customer.billState, customer.billZip)
    phone = customer.fullName
    email = customer.email
    p = document.add_paragraph()

    # Create table
    table = document.add_table(rows=3, cols=2)

    # fill in row 1
    row1 = table.rows[0].cells
    row1[0].text = '%s' % (customer.fullName,)
    row1[1].text = '%s' % (today,)
    row1[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # fill in row 2
    row2 = table.rows[1].cells
    row2[0].text = bill1
    row2[1].text = customer.phone1
    row2[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # fill in row 3
    row3 = table.rows[2].cells
    row3[0].text = bill2
    row3[1].text = customer.email
    row3[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Change Font size of Table
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10)


    # add Job location
    p = document.add_paragraph()
    p.paragraph_format.line_spacing = Pt(blank_space)
    p = document.add_paragraph()
    runner = p.add_run('JOB LOCATION: %s' % (jobsite.jobStreet,))
    runner.font.size = Pt(10)
    runner.font.bold = True
    runner.font.all_caps = True
    # Add invoice line
    p = document.add_paragraph()
    runner = p.add_run('**************************************************INVOICE************************************************')
    runner.font.bold = True
    runner.font.size = Pt(10)
    runner.font.all_caps = True
    runner.font.highlight_color = WD_COLOR_INDEX.YELLOW

    sections = estimate.section_set.all()
    for section in sections:
        p = document.add_paragraph()
        p.paragraph_format.line_spacing = Pt(blank_space)

        #Add heading for each section
        p= document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.keep_together = True
        runner = p.add_run(section.heading +' ('+'${:,.2f}'.format(section.price) +')')
        runner.font.size = Pt(10)
        runner.font.all_caps = True
        runner.bold = True
        runner.underline = True

        html = section.description
        text = html2text.html2text(html)


        spl = text.split('\n')

        # Sanity
        print(spl)
        i = 0
        for i in range(0, len(spl)):
            sect = spl[i]
            p = document.add_paragraph()
            p.paragraph_format.keep_together = True
            p.paragraph_format.keep_with_next = True
            runner = p.add_run()
            runner.add_tab()
            runner = p.add_run(sect)
            runner.font.size = Pt(10)



        runner = p.add_run(section.get_guarantee_display())
        runner.font.size = Pt(10)
        runner.font.bold = True
        runner.font.underline = True

        # p.alignment = WD_ALIGN_PARAGRAPH.LEFT


    p = document.add_paragraph()
    runner = p.add_run()
    runner.add_tab()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    total = '${:,.2f}'.format(estimate.total)
    runner = p.add_run('BALANCE DUE: %s' % (total,))
    runner.font.size = Pt(10)
    runner.font.bold = True
    runner.font.highlight_color = WD_COLOR_INDEX.YELLOW


    # add job complete
    p = document.add_paragraph()
    dt = estimate.completedDate
    runner = p.add_run()
    runner.add_tab()
    runner = p.add_run('JOB COMPLETE %s' % (dt.strftime('%m/%d/%y')))
    runner.font.color.rgb = RGBColor(150, 0, 0)
    runner.font.bold = True
    runner.font.size = Pt(10)

    # p = document.add_paragraph()
    # runner = p.add_run()
    # runner.add_tab()
    # runner = p.add_run('BALANCE DUE')
    # runner.font.color.rgb = RGBColor(150, 0, 0)
    # runner.font.bold = True
    # runner.font.size = Pt(10)
    #
    # p = document.add_paragraph()
    # runner = p.add_run()
    # runner.add_tab()
    # runner = p.add_run('THANK YOU')
    # runner.font.color.rgb = RGBColor(150, 0, 0)
    # runner.font.size = Pt(10)
    # runner.font.bold = True



    # add thank you note
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(5)
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = True
    runner = p.add_run('THANK YOU')
    runner.font.size = Pt(10)
    runner.font.bold = True

    # add user name
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = True
    user = request.user
    username = '%s %s' % (user.first_name, user.last_name)
    username = username.upper()
    p.paragraph_format.left_indent = Inches(5)
    runner = p.add_run('%s' % (username,))  # add in username from logged in user
    runner.font.size = Pt(10)
    runner.font.bold = True

    # add reiter roofing
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(5)
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.widow_control = True
    runner = p.add_run('REITER ROOFING')
    runner.font.size = Pt(10)
    runner.font.bold = True

    # save docx
    document.save(response)

    # Get the value of the BytesIO buffer and write it to the response.
    docx = buffer.getvalue()
    buffer.close()
    response.write(docx)
    return response




def write_docx_view(request, cust, job, ticket, est):
    import html2text
    import re
    from docx.shared import RGBColor, Pt, Inches
    from docx.enum.text import WD_COLOR_INDEX

    current_path = request.get_full_path()
    print(current_path)

    # load objects
    customer = get_object_or_404(Customer, id=cust)
    jobsite = get_object_or_404(Jobsite, id=job)
    ticket = get_object_or_404(Ticket, id=ticket)
    estimate = get_object_or_404(Estimate, id=est)

    estimate.completed = True
    estimate.save()

    # setup buffer
    buffer = BytesIO()

    # Set Some Variables for filename and path
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    # headerimage = os.path.join(BASE_DIR, "static/CRM/img/logo-redux.png")
    filename = customer.lastName + "_" + customer.firstName + "-" + jobsite.jobStreet + "-workorder#" + str(ticket.id)
    doc = os.path.join(BASE_DIR, "static/CRM/doc-template/newest.docx")

    # Set misc variables
    blank_space = 6
    topmargin = .50
    gutters = .50
    bottommargin = .25

    # Create HttpResponse object with appropriate PDF headers
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachement; filename=' + filename + ".docx"

    # Create the docx object
    document = Document(doc)

    sections = document.sections
    for section in sections:
        section.top_margin = Inches(topmargin)
        section.bottom_margin = Inches(bottommargin)
        section.left_margin = Inches(gutters)
        section.right_margin = Inches(gutters)

    # Set the filename variable
    docx_title = "%s.docx" % (filename,)

    # Add content to docx file

    # Add blank line
    p = document.add_paragraph()


    # setup variables for docx table
    today = datetime.date.today()
    bill1 = customer.billStreet
    bill2 = '%s, %s, %s' % (customer.billCity, customer.billState, customer.billZip)
    phone = customer.fullName
    email = customer.email

    # Create table
    table = document.add_table(rows=3, cols=2)

    # fill in row 1
    row1 = table.rows[0].cells
    row1[0].text = '%s' % (customer.fullName,)
    row1[1].text = '%s' % (today,)
    row1[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # fill in row 2
    row2 = table.rows[1].cells
    row2[0].text = bill1
    row2[1].text = customer.phone1
    row2[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # fill in row 3
    row3 = table.rows[2].cells
    row3[0].text = bill2
    row3[1].text = customer.email
    row3[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Change Font size of Table
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10)

    # add Job location
    p = document.add_paragraph()
    p.paragraph_format.line_spacing = Pt(blank_space)
    p = document.add_paragraph()
    runner = p.add_run('JOB LOCATION: %s' % (jobsite.jobStreet,))
    runner.font.size = Pt(10)
    runner.font.bold = True
    runner.font.all_caps = True

    # Add Estimate Line
    p=document.add_paragraph()

    runner = p.add_run(
        '**************************************************ESTIMATE***********************************************')
    runner.font.bold = True
    runner.font.size = Pt(10)
    runner.font.all_caps = True
    runner.font.highlight_color = WD_COLOR_INDEX.YELLOW

    sections = estimate.section_set.all()
    for section in sections:
        p = document.add_paragraph()
        p.paragraph_format.line_spacing = Pt(blank_space)

        # Add heading for each section
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.keep_together = True
        runner = p.add_run(section.heading + ' (' + '${:,.2f}'.format(section.price) + ')')
        runner.font.size = Pt(10)
        runner.font.all_caps = True
        runner.bold = True
        runner.underline = True

        html = section.description
        text = html2text.html2text(html)

        spl = text.split('\n')

        # Sanity
        print(spl)
        i = 0
        for i in range(0, len(spl)):
            sect = spl[i]
            p = document.add_paragraph()
            p.paragraph_format.keep_together = True
            p.paragraph_format.keep_with_next = True
            runner = p.add_run()
            runner.add_tab()
            runner = p.add_run(sect)
            runner.font.size = Pt(10)

        runner = p.add_run(section.get_guarantee_display())
        runner.font.size = Pt(10)
        runner.font.bold = True
        runner.font.underline = True



    p = document.add_paragraph()
    runner = p.add_run()
    runner.add_tab()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    total = '${:,.2f}'.format(estimate.total)
    runner = p.add_run('TOTAL PRICE: %s' % (total,))
    runner.font.size = Pt(10)
    runner.font.bold = True
    # runner.font.highlight_color = WD_COLOR_INDEX.YELLOW




    # add thank you note
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(5)
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = True
    runner = p.add_run('THANK YOU')
    runner.font.size = Pt(10)
    runner.font.bold = True

    # add user name
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = True
    user = request.user
    username = '%s %s' % (user.first_name, user.last_name)
    username = username.upper()
    p.paragraph_format.left_indent = Inches(5)
    runner = p.add_run('%s' % (username,))  # add in username from logged in user
    runner.font.size = Pt(10)
    runner.font.bold = True

    # add reiter roofing
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(5)
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.widow_control = True
    runner = p.add_run('REITER ROOFING')
    runner.font.size = Pt(10)
    runner.font.bold = True

    # save docx
    document.save(response)

    # Get the value of the BytesIO buffer and write it to the response.
    docx = buffer.getvalue()
    buffer.close()
    response.write(docx)
    return response


def write_pdf_view(request, cust, job, ticket, est):
    current_path = request.get_full_path()
    print(current_path)

    customer = get_object_or_404(Customer, id=cust)
    jobsite = get_object_or_404(Jobsite, id=job)
    ticket = get_object_or_404(Ticket, id=ticket)

    filename = customer.lastName + "_" + customer.firstName + "-" + jobsite.jobStreet + "-workorder#" + str(ticket.id)

    # Create HttpResponse object with appropriate PDF headers
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachement; filename=' + filename + ".pdf"

    # Create the PDF object, using the response object as its "file"
    p = canvas.Canvas(response, pagesize=letter)
    width, height = letter

    # Setup buffer
    buffer = BytesIO()

    bill_add2 = customer.billCity + " " + customer.billState + "," + customer.billZip
    job_add2 = jobsite.jobCity + " " + jobsite.jobState + "," + jobsite.jobZip

    created = str(ticket.created)
    created = created[:10]

    # Draw things on the PDF. Here's where the PDF generation happens.
    # See the ReportLab documentation for the full list of functionality.
    p.setFont("Times-Bold", 12)
    p.drawString(50, height - 51, "Customer Info:")
    p.setFont("Times-Bold", 20)
    p.drawString(225, height - 25, "WORK ORDER")
    p.line(25, height - 27, width - 25, height - 27)

    p.setFont("Times-Roman", 12)

    p.drawString(50, height - 65, customer.fullName)
    p.drawString(50, height - 79, customer.billStreet)
    p.drawString(50, height - 93, bill_add2)
    p.drawString(50, height - 107, customer.phone1)
    p.drawString(50, height - 121, customer.phone2)
    p.drawString(50, height - 135, customer.email)
    p.drawString(50, height - 149, 'Source: ' + customer.source)

    p.setFont("Times-Bold", 12)
    p.drawString(225, height - 51, "Jobsite Info:")

    p.setFont("Times-Roman", 12)
    p.drawString(225, height - 65, jobsite.jobStreet)
    p.drawString(225, height - 79, job_add2)
    p.drawString(225, height - 93, "Stories: " + str(jobsite.stories))
    p.drawString(225, height - 107, "Access: " + jobsite.access)

    p.setFont("Times-Bold", 12)
    p.drawString(50, height - 172, "Workorder Info:")

    p.setFont("Times-Roman", 12)

    # make neon yellow box
    p.setFillColorRGB(204, 255, 0, alpha=None)
    p.rect(398, height - 69, 100, 16, stroke=1, fill=1)

    # set color back to black for text
    p.setFillColorRGB(0, 0, 0, alpha=None)

    p.drawString(400, height - 65, "Call type: " + ticket.call_type)
    p.drawString(400, height - 51, "Date created: " + created)
    p.drawString(50, height - 186, "Problem: " + ticket.problem)
    p.drawString(50, height - 200, "Notes: " + ticket.notes)

    p.line(25, height - 227, width - 25, height - 227)

    # Close the PDF object cleanly, and we're done.
    p.showPage()
    p.save()

    # Get the value of the BytesIO buffer and write it to the response.
    pdf = buffer.getvalue()
    buffer.close()
    response.write(pdf)
    return response
