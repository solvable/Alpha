import datetime
import os
from io import BytesIO

from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.urls import reverse, reverse_lazy
from django.views import generic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from .models import Customer, Jobsite, Ticket


def write_docx_view(request, cust, job, ticket):
    current_path = request.get_full_path()
    print(current_path)

    # load objects
    customer = get_object_or_404(Customer, id=cust)
    jobsite = get_object_or_404(Jobsite, id=job)
    ticket = get_object_or_404(Ticket, id=ticket)

    # setup buffer
    buffer = BytesIO()

    # Set Some Variables for filename and path
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    headerimage = os.path.join(BASE_DIR, "static/CRM/img/logo-redux.png")
    filename = customer.lastName + "_" + customer.firstName + "-" + jobsite.jobStreet + "-workorder#" + str(ticket.id)
    doc = os.path.join(BASE_DIR, "static/CRM/doc-template/newest.docx")

    # Create HttpResponse object with appropriate PDF headers
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachement; filename=' + filename + ".docx"

    # Create the docx object
    document = Document(doc)

    # Set the filename variable
    docx_title = "%s.docx" % (filename,)

    # Add content to docx file

    # setup variables for docx table
    today = datetime.date.today()
    bill1 = customer.billStreet
    bill2 = '%s, %s, %s' % (customer.billCity, customer.billState, customer.billZip)
    phone = customer.fullName
    email = customer.email

    # Create table
    table = document.add_table(rows=3, cols=2)

    # fill in row 1
    row1 = table.rows[0].cells
    row1[0].text = '%s' % (customer.fullName,)
    row1[1].text = '%s' % (today,)
    row1[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # fill in row 2
    row2 = table.rows[1].cells
    row2[0].text = bill1
    row2[1].text = customer.phone1
    row2[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # fill in row 3
    row3 = table.rows[2].cells
    row3[0].text = bill2
    row3[1].text = customer.email
    row3[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # add Job location
    p = document.add_paragraph()
    p = document.add_paragraph()
    p.add_run('JOB LOCATION: %s' % (jobsite.jobStreet,)).bold = True

    p = document.add_paragraph()
    p = document.add_paragraph()
    p = document.add_paragraph()
    p = document.add_paragraph()
    p = document.add_paragraph()
    p = document.add_paragraph()
    p = document.add_paragraph()

    p.add_run('TOTAL PRICE: $').bold = True
    p = document.add_paragraph()

    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(4.5)
    p.add_run('THANK YOU').bold = True
    # p = document.add_paragraph()
    # p.paragraph_format.left_indent = Inches(0.5)
    # p.add_run('USERNAME').bold = True // add in username from logged in user
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(4.5)
    p.add_run('REITER ROOFING').bold = True

    # save docx
    document.save(response)

    # Get the value of the BytesIO buffer and write it to the response.
    docx = buffer.getvalue()
    buffer.close()
    response.write(docx)
    return response


def write_pdf_view(request, cust, job, ticket):
    current_path = request.get_full_path()
    print(current_path)

    customer = get_object_or_404(Customer, id=cust)
    jobsite = get_object_or_404(Jobsite, id=job)
    ticket = get_object_or_404(Ticket, id=ticket)

    filename = customer.lastName + "_" + customer.firstName + "-" + jobsite.jobStreet + "-workorder#" + str(ticket.id)

    # Create HttpResponse object with appropriate PDF headers
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachement; filename=' + filename + ".pdf"

    # Create the PDF object, using the response object as its "file"
    p = canvas.Canvas(response, pagesize=letter)
    width, height = letter

    # Setup buffer
    buffer = BytesIO()

    bill_add2 = customer.billCity + " " + customer.billState + "," + customer.billZip
    job_add2 = jobsite.jobCity + " " + jobsite.jobState + "," + jobsite.jobZip

    created = str(ticket.created)
    created = created[:10]

    # Draw things on the PDF. Here's where the PDF generation happens.
    # See the ReportLab documentation for the full list of functionality.
    p.setFont("Times-Bold", 12)
    p.drawString(50, height - 51, "Customer Info:")
    p.setFont("Times-Bold", 20)
    p.drawString(225, height - 25, "WORK ORDER")
    p.line(25, height - 27, width - 25, height - 27)

    p.setFont("Times-Roman", 12)

    p.drawString(50, height - 65, customer.fullName)
    p.drawString(50, height - 79, customer.billStreet)
    p.drawString(50, height - 93, bill_add2)
    p.drawString(50, height - 107, customer.phone1)
    p.drawString(50, height - 121, customer.phone2)
    p.drawString(50, height - 135, customer.email)
    p.drawString(50, height - 149, 'Source: ' + customer.source)

    p.setFont("Times-Bold", 12)
    p.drawString(225, height - 51, "Jobsite Info:")

    p.setFont("Times-Roman", 12)
    p.drawString(225, height - 65, jobsite.jobStreet)
    p.drawString(225, height - 79, job_add2)
    p.drawString(225, height - 93, "Stories: " + str(jobsite.stories))
    p.drawString(225, height - 107, "Access: " + jobsite.access)

    p.setFont("Times-Bold", 12)
    p.drawString(50, height - 172, "Workorder Info:")

    p.setFont("Times-Roman", 12)

    # make neon yellow box
    p.setFillColorRGB(204, 255, 0, alpha=None)
    p.rect(398, height - 69, 100, 16, stroke=1, fill=1)

    # set color back to black for text
    p.setFillColorRGB(0, 0, 0, alpha=None)

    p.drawString(400, height - 65, "Call type: " + ticket.call_type)
    p.drawString(400, height - 51, "Date created: " + created)
    p.drawString(50, height - 186, "Problem: " + ticket.problem)
    p.drawString(50, height - 200, "Notes: " + ticket.notes)

    p.line(25, height - 227, width - 25, height - 227)

    # Close the PDF object cleanly, and we're done.
    p.showPage()
    p.save()

    # Get the value of the BytesIO buffer and write it to the response.
    pdf = buffer.getvalue()
    buffer.close()
    response.write(pdf)
    return response


class IndexView(generic.ListView):
    model = Ticket
    template_name = 'CRM/index.html'

    latest_tickets = Ticket.objects.order_by('-created')[:10]
    open_tickets = Ticket.objects.filter(completed=False)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['latest_tickets'] = IndexView.latest_tickets
        context['open_tickets'] = IndexView.open_tickets
        return context


'''
Customer Views
'''


class CustomerDetailView(generic.DetailView):
    model = Customer
    template_name = 'CRM/customer_detail.html'
    pk_url_kwarg = "cust"


class CustomerCreateView(generic.CreateView):
    model = Customer
    template_name = 'CRM/customer_create.html'
    pk_url_kwarg = "cust"
    fields = ['firstName', 'lastName', 'billStreet', 'billCity', 'billState', 'billZip', 'phone1', 'phone2', 'email',
              'source']


class CustomerUpdateView(generic.UpdateView):
    model = Customer
    template_name = 'CRM/customer_update.html'
    pk_url_kwarg = "cust"
    fields = ['firstName', 'lastName', 'billStreet', 'billCity', 'billState', 'billZip', 'phone1', 'phone2', 'email',
              'source']


class CustomerDeleteView(generic.DeleteView):
    model = Customer
    pk_url_kwarg = 'cust'
    success_url = reverse_lazy('index')


'''
Jobsite Views
'''


class JobsiteDetailView(generic.DetailView):
    model = Jobsite
    template_name = 'CRM/jobsite_detail.html'
    pk_url_kwarg = "job"




    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['customer'] = get_object_or_404(Customer, id=self.kwargs['cust'])
        return context


class JobsiteCreateView(generic.CreateView):
    model = Jobsite
    template_name = 'CRM/jobsite_create.html'
    pk_url_kwarg = "cust"
    fields = ['customer_id', 'jobStreet', 'jobCity', 'jobState', 'jobZip', 'stories', 'access', 'notes']

    def get_initial(self):
        customer = get_object_or_404(Customer, pk=self.kwargs.get('cust'))
        return {
            'customer_id': customer.id,
        }


class JobsiteUpdateView(generic.UpdateView):
    model = Jobsite
    template_name = 'CRM/jobsite_create.html'
    fields = ['customer_id', 'jobStreet', 'jobCity', 'jobState', 'jobZip', 'stories', 'access', 'notes']

    pk_url_kwarg = "job"


class JobsiteDeleteView(generic.DeleteView):
    model = Jobsite
    pk_url_kwarg = "job"

    # template_name = 'CRM/jobsite_delete.html'
    def get_success_url(self):
        return reverse('customer_detail', kwargs={'cust': self.kwargs.get('cust')})


'''
Ticket Views
'''


class TicketDetailView(generic.DetailView):
    model = Ticket
    template_name = 'CRM/ticket_detail.html'
    pk_url_kwarg = "ticket"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['customer'] = get_object_or_404(Customer, id=self.kwargs['cust'])
        context['jobsite'] = get_object_or_404(Jobsite, id=self.kwargs['job'])
        return context


class TicketCreateView(generic.CreateView):
    model = Ticket
    template_name = 'CRM/ticket_create.html'
    pk_url_kwarg = "job"
    fields = ['customer_id', 'jobsite_id', 'call_type', 'problem', 'completed', 'notes']

    def get_initial(self):
        customer = get_object_or_404(Customer, id=self.kwargs.get('cust'))
        jobsite = get_object_or_404(Jobsite, id=self.kwargs.get('job'))

        return {
            'customer_id': customer.id,
            'jobsite_id': jobsite.id
        }


class TicketUpdateView(generic.UpdateView):
    model = Ticket
    template_name = 'CRM/ticket_create.html'
    fields = ['customer_id', 'jobsite_id', 'call_type', 'problem', 'completed', 'notes']
    pk_url_kwarg = "ticket"


class TicketDeleteView(generic.DeleteView):
    model = Ticket
    pk_url_kwarg = "ticket"

    def get_success_url(self):
        return reverse('jobsite_detail', kwargs={'cust': self.kwargs.get('cust'), 'job': self.kwargs.get('job')})
